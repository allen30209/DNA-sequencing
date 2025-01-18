import os
from docx import Document


def find_dna_sequences(dna_sequence, start_codon, stop_codon, min_length=400, max_length=500):
    """
    找出 DNA 序列中位於指定起始密碼和終止密碼之間的完整序列，並限制長度範圍。
    """
    results = []
    start_idx = 0

    while start_idx < len(dna_sequence):
        # 找起始密碼
        start_idx = dna_sequence.find(start_codon, start_idx)
        if start_idx == -1:
            break  # 找不到更多起始密碼，結束

        # 找終止密碼
        stop_idx = dna_sequence.find(stop_codon, start_idx + len(start_codon))
        if stop_idx == -1:
            break  # 找不到終止密碼，結束

        # 提取完整序列（包含起始和終止密碼）
        full_sequence = dna_sequence[start_idx:stop_idx + len(stop_codon)]
        sequence_length = len(full_sequence)

        # 確認序列大小是否在範圍內
        if min_length <= sequence_length <= max_length:
            results.append((full_sequence, sequence_length))

        # 更新索引，避免死循環
        start_idx = stop_idx + len(stop_codon)

    return results


def process_dna_folder(input_folder , output_folder, start_codon, stop_codon):
    """
    讀取資料夾中的 DNA 檔案，進行分析並輸出結果到新資料夾和 Word 文件中。
    """
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # 建立 Word 文件
    doc = Document()
    doc.add_heading("DNA 分析結果", level=1)

    for filename in os.listdir(input_folder):
        if filename.endswith(".txt"):
            file_path = os.path.join(input_folder, filename)

            # 讀取 DNA 序列檔案
            with open(file_path, "r") as file:
                dna_sequence = file.read().replace(" ", "").replace("\n", "").strip()

            # 分析 DNA
            results = find_dna_sequences(dna_sequence, start_codon, stop_codon)

            # 結果輸出到文字檔案
            result_file_path = os.path.join(output_folder, f"result_{filename}")
            with open(result_file_path, "w") as result_file:
                if results:
                    for i, (seq, length) in enumerate(results, 1):
                        result_file.write(f"結果 {i}:\n")
                        result_file.write(f"序列長度: {length}\n")
                        result_file.write(f"序列:\n{seq}\n\n")
                else:
                    result_file.write("未找到符合條件的序列。\n")

            # 匯總到 Word 文件
            doc.add_heading(f"檔案: {filename}", level=2)
            if results:
                for i, (seq, length) in enumerate(results, 1):
                    doc.add_paragraph(f"結果 {i}：")
                    doc.add_paragraph(f"序列長度: {length}")
                    doc.add_paragraph(f"序列: {seq}")
            else:
                doc.add_paragraph("未找到符合條件的序列。")

    # 儲存 Word 文件
    word_file_path = os.path.join(output_folder, "DNA_Analysis_Results.docx")
    doc.save(word_file_path)


# 使用範例
input_folder = "input_dna"  # 輸入資料夾路徑
output_folder = "output_results"  # 輸出資料夾路徑
start_codon = ""  # 起始密碼
stop_codon = ""  # 終止密碼

# 執行處理
process_dna_folder(input_folder, output_folder, start_codon, stop_codon)